"""
報道機関対応報告書テンプレート（.docx）生成スクリプト

nested_form_builder の {@field|pipe} トークンを埋め込んだ Word テンプレートを生成する。
ユーザーは生成された .docx を Google Drive にアップロード → Google Doc 化 →
そのURLを ConfigPage の「標準印刷テンプレートURL」欄に貼り付けて運用する。

使い方:
    python scripts/generate_press_template.py
"""

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\sa11882\Downloads\報道対応報告書テンプレート.docx"
JP_FONT = "MS Mincho"  # 游明朝 / MS明朝。Google Docへ変換時にも安定


def set_cell_borders(cell):
    """セル四方に黒い細罫線を引く"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_shading(cell, fill_hex):
    """セル背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_width(cell, mm):
    """セル幅をミリ指定"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(mm * 56.7)))  # 1mm ≒ 56.7 twips
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_run_font(run, size_pt=10.5, bold=False):
    run.font.name = JP_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), JP_FONT)
    rFonts.set(qn("w:ascii"), JP_FONT)
    rFonts.set(qn("w:hAnsi"), JP_FONT)


def write_cell(cell, lines, *, bold=False, size_pt=10.5,
               align=WD_ALIGN_PARAGRAPH.LEFT, vertical=WD_ALIGN_VERTICAL.CENTER,
               shading=None):
    """セルに複数行テキストを書き込む。リスト各要素が1段落となる。"""
    cell.text = ""  # clear default empty paragraph
    cell.vertical_alignment = vertical
    if shading:
        set_cell_shading(cell, shading)
    set_cell_borders(cell)
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = p.add_run(line)
        set_run_font(run, size_pt=size_pt, bold=bold)


def merge_vertical(table, col, start_row, end_row):
    """同一列で複数行をvMerge"""
    a = table.cell(start_row, col)
    for r in range(start_row + 1, end_row + 1):
        a = a.merge(table.cell(r, col))


def main():
    doc = Document()

    # ページ設定（A4・余白20mm）
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # ===== タイトル =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("報道機関からの取材・問い合わせへの対応報告書")
    set_run_font(run, size_pt=14, bold=True)

    # ===== ヘッダ行（日付・所属・氏名・電話） =====
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(
        "{@対応日|time:gge年M月D日}　　所属　{@報告者所属}　氏名　{@報告者名}　℡　{@報告者電話番号}"
    )
    set_run_font(run, size_pt=10.5)

    # ===== 表本体 =====
    # 行構成:
    #   行0: 内容区分(label)         | 件名+チェック
    #   行1: 相手方区分(label,2行縦結合) | 区分(label) | チェック
    #   行2:                         | 社名(label)  | 社名+番組名
    #   行3: 取材目的                | 取材目的本文
    #   行4: 対応日時方法 〔対応者〕  | 日時+方法+対応者
    #   行5: 取材内容/今後...         | 本文
    #   行6: ※後日取材の場合 取材内容 | 空欄
    #   行7: 報道の結果              | チェック群 | 備考

    # 列幅は3列：30mm / 25mm / 残り
    LBL_W = 32   # 左ラベル列
    SUB_W = 22   # サブラベル列（行1, 2のみ使用）
    RIGHT_W = 110  # 値列
    BIKO_W = 50  # 行7の右端「備考」列

    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    SHD = "F2F2F2"  # ラベル背景

    # --- Row 0: 内容区分 ---
    row = table.rows[0]
    # ラベル(0,0)+(0,1)を結合して1セル化
    c0 = row.cells[0].merge(row.cells[1])
    write_cell(c0, ["内容", "区分"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(c0, LBL_W + SUB_W)
    write_cell(
        row.cells[2],
        [
            "【件名】{@件名}",
            "　　　　　　　　　{@内容区分|ifv:@内容区分==取材,■,□}取材　　{@内容区分|ifv:@内容区分==取材準備,■,□}取材準備　　{@内容区分|ifv:@内容区分==問い合わせ,■,□}問い合わせ",
        ],
    )
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 1: 相手方 / 区分 / チェック ---
    row = table.rows[1]
    write_cell(row.cells[0], ["相手方"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(row.cells[0], LBL_W)
    write_cell(row.cells[1], ["区分"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(row.cells[1], SUB_W)
    write_cell(
        row.cells[2],
        [
            "{@相手方区分|ifv:@相手方区分==新聞社,■,□}新聞社　{@相手方区分|ifv:@相手方区分==テレビ局,■,□}テレビ局　{@相手方区分|ifv:@相手方区分==ラジオ局,■,□}ラジオ局　{@相手方区分|ifv:@相手方区分==雑誌社,■,□}雑誌社　{@相手方区分|ifv:@相手方区分==その他,■,□}その他（{@相手方区分（その他）}）",
        ],
    )
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 2: (相手方vMerge) / 社名 / 社名+番組名 ---
    row = table.rows[2]
    write_cell(row.cells[1], ["社名"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(row.cells[1], SUB_W)
    write_cell(
        row.cells[2],
        ["【社名(担当)】{@社名}　{@取材担当者}　　　番組名　{@番組名}"],
    )
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)
    # 行1と行2の左端「相手方」セルを縦結合
    merge_vertical(table, 0, 1, 2)
    # 結合後の左セルに「相手方」表示はrow1で書き済み

    # --- Row 3: 取材目的 ---
    row = table.rows[3]
    c = row.cells[0].merge(row.cells[1])
    write_cell(c, ["取材目的"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(c, LBL_W + SUB_W)
    write_cell(row.cells[2], ["{@取材目的}"])
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 4: 対応日時方法 〔対応者〕 ---
    row = table.rows[4]
    c = row.cells[0].merge(row.cells[1])
    write_cell(c, ["対応日時・方法", "〔対応者〕"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(c, LBL_W + SUB_W)
    write_cell(
        row.cells[2],
        [
            "{@対応日|time:M月D日}({@対応日|time:ddd}){@対応開始時間}～{@対応終了時間}",
            "{@対応方法|ifv:@対応方法==来庁,■,□}来庁　　{@対応方法|ifv:@対応方法==電話,■,□}電話　　{@対応方法|ifv:@対応方法==その他,■,□}その他（{@対応方法（その他）}）　　対応：{@報告者名}",
        ],
    )
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 5: 取材内容及びその対応 / 今後の取材日程 ---
    row = table.rows[5]
    c = row.cells[0].merge(row.cells[1])
    write_cell(c, ["取材内容", "及びその対応", "", "今後の取材日程・", "内容・対応", "予定者など"],
               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(c, LBL_W + SUB_W)
    write_cell(
        row.cells[2],
        [
            "{@取材内容及びその対応}",
            "",
            "──────────",
            "【今後の予定】",
            "{@今後の取材日程・内容・対応予定者など}",
        ],
    )
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 6: ※後日取材の場合 取材内容 ---
    row = table.rows[6]
    c = row.cells[0].merge(row.cells[1])
    write_cell(c, ["※後日取材の場合", "取材内容"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(c, LBL_W + SUB_W)
    write_cell(row.cells[2], [""])
    set_cell_width(row.cells[2], RIGHT_W + BIKO_W)

    # --- Row 7: 報道の結果 / 備考 ---
    # 既定の3列構成をそのまま利用（label / 値 / 備考）
    row = table.rows[7]
    write_cell(row.cells[0], ["報道の結果"], bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER, shading=SHD)
    set_cell_width(row.cells[0], LBL_W)
    write_cell(
        row.cells[1],
        [
            "{@報道の結果|ifv:記事掲載 in _,■,□}記事掲載（{@記事掲載日|time:M/D}）",
            "{@報道の結果|ifv:放送予定 in _,■,□}放送予定（{@放送予定詳細}）",
            "{@報道の結果|ifv:その他 in _,■,□}その他（{@報道結果（その他）}）",
            "{@報道の結果|ifv:報道なし in _,■,□}報道なし",
        ],
    )
    set_cell_width(row.cells[1], SUB_W + RIGHT_W)
    write_cell(
        row.cells[2],
        ["【備考】", "{@備考}"],
    )
    set_cell_width(row.cells[2], BIKO_W)

    # ===== 出力 =====
    doc.save(OUT_PATH)
    print(f"OK: {OUT_PATH}")


if __name__ == "__main__":
    main()
